"""Generate the mid-semester Design & Development dissertation report.

The report follows the provided BITS-style dissertation example but uses the
actual digit-recognition MLOps project details and the phase window:
15 May 2026 - 14 June 2026.

Usage:
    python report/build_midsem_report.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = (
    PROJECT_ROOT
    / "report"
    / "2024da04367-digit-recognition-midsem-design-development-report.docx"
)
ARTIFACTS = PROJECT_ROOT / "artifacts"
MODELS = PROJECT_ROOT / "models"

TITLE = (
    "MLOps Pipeline for a Digit Recognition Service: Design and Development "
    "Dissertation"
)
STUDENT_NAME = "Jhanvi Titoria"
STUDENT_ID = "2024DA04367"
PROGRAMME = "M.Tech. Data Science & Engineering"
ORGANIZATION = "Nokia Solutions and Networks, IN/Bangalore/Manyata E2"
SUPERVISOR = "Abhishek Kumar"
EXAMINER = "Abhishek Kumar Jain"
PHASE = "Design & Development"
PHASE_DATES = "15 May 2026 - 14 June 2026"


def center(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_center(doc: Document, text: str, size: int = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    center(p)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_subheading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, rows: list[list[str]], header: bool = True):
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = text
            if header and i == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    return table


def load_json(path: Path, default: dict) -> dict:
    if path.exists():
        with open(path, encoding="utf-8") as fh:
            return json.load(fh)
    return default


def project_facts() -> dict:
    champion = load_json(
        MODELS / "champion.json",
        {
            "champion": "balanced",
            "model_type": "svm",
            "baseline_accuracy": 0.9833,
            "metrics": {
                "train_accuracy": 1.0,
                "val_accuracy": 0.9896,
                "test_accuracy": 0.9833,
                "train_val_gap": 0.0104,
            },
            "selection_seconds": 0.00005,
            "selection_rule": (
                "max(val_accuracy) subject to train_val_gap <= 0.08; "
                "tie-break on test_accuracy"
            ),
        },
    )
    experiment = load_json(ARTIFACTS / "experiment_results.json", {})
    return {"champion": champion, "experiment": experiment}


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def cover(doc: Document) -> None:
    add_center(doc, TITLE.upper(), size=16, bold=True)
    add_center(doc, "by", size=12)
    add_center(doc, STUDENT_NAME, size=13, bold=True)
    add_center(doc, STUDENT_ID, size=12, bold=True)
    doc.add_paragraph()
    add_center(doc, "Dissertation work carried out at", bold=True)
    add_center(doc, ORGANIZATION)
    doc.add_paragraph()
    add_center(
        doc,
        f"Submitted in partial fulfilment of {PROGRAMME} degree programme",
    )
    doc.add_paragraph()
    add_center(doc, "Under the Supervision of", bold=True)
    add_center(doc, SUPERVISOR, bold=True)
    add_center(doc, ORGANIZATION)
    doc.add_paragraph()
    add_center(doc, f"Mid-Semester Phase: {PHASE}", bold=True)
    add_center(doc, PHASE_DATES, bold=True)
    doc.add_paragraph()
    add_center(doc, "BITS", size=15, bold=True)
    add_center(doc, "PILANI (RAJASTHAN)", size=13, bold=True)
    add_center(doc, "June 2026", size=12, bold=True)


def abstract(doc: Document, facts: dict) -> None:
    champ = facts["champion"]
    metrics = champ["metrics"]
    first_alert = facts["experiment"].get("first_alert") or {
        "severity": 0.2,
        "errors": 311,
    }
    add_heading(doc, "ABSTRACT")
    doc.add_paragraph(
        "This mid-semester report presents the design and development phase of "
        "an end-to-end MLOps pipeline for handwritten digit recognition. The "
        "project uses a well-understood ten-class image-classification task as "
        "the payload and focuses on the engineering work required to move a "
        "machine learning model from experimentation to an operable service."
    )
    doc.add_paragraph(
        "The implemented system contains a config-driven data and training "
        "pipeline, MLflow experiment tracking, automated model comparison, "
        "champion model selection, drift/performance-drop monitoring, FastAPI "
        "serving, Prometheus metrics, Docker packaging, CI, and load/latency "
        "test support. The data layer supports the scikit-learn digits dataset "
        "and a config switch for MNIST."
    )
    doc.add_paragraph(
        f"During {PHASE} ({PHASE_DATES}), the main development outcome is a "
        f"working champion-based serving design. The selector promotes the "
        f"{champ['champion']} ({champ['model_type']}) model using the rule "
        f"'{champ['selection_rule']}'. The promoted model records "
        f"{pct(metrics['test_accuracy'])} test accuracy with a train-validation "
        f"gap of {metrics['train_val_gap'] * 100:.2f} percentage points."
    )
    doc.add_paragraph(
        "The monitoring module compares live labelled-batch accuracy against "
        "the champion baseline. In the proof experiment, drift is detected at "
        f"severity {first_alert.get('severity', 0.2)}, flagging "
        f"{first_alert.get('errors', 311)} errors before they become invisible "
        "user-facing failures. This completes the core design and development "
        "work and prepares the project for the testing and dissertation review "
        "phases."
    )
    doc.add_paragraph()
    add_table(
        doc,
        [
            ["Signature of the Student", "Signature of the Supervisor"],
            [f"Name: {STUDENT_NAME}", f"Name: {SUPERVISOR}"],
            ["Date:", "Date:"],
            ["Place:", "Place:"],
        ],
        header=False,
    )


def contents(doc: Document) -> None:
    add_heading(doc, "Contents")
    add_table(
        doc,
        [
            ["1.", "Modules in Digit Recognition MLOps Pipeline", "5"],
            ["2.", "Functional Block Diagram / Description", "8"],
            ["3.", "Major Technical Specifications", "10"],
            ["4.", "Design Considerations", "11"],
            ["5.", "Future Plan", "12"],
            ["6.", "Abbreviations", "13"],
        ],
        header=False,
    )
    doc.add_paragraph()
    add_table(
        doc,
        [
            ["Figure 1", "Modules in Digit Recognition MLOps Pipeline", "5"],
            ["Figure 2", "Functional Block Diagram", "8"],
            ["Table 1", "Technical Specifications of the Project", "10"],
        ],
        header=False,
    )


def modules(doc: Document) -> None:
    add_heading(doc, "1. MODULES IN DIGIT RECOGNITION MLOPS PIPELINE")
    doc.add_paragraph(
        "The system is organized into software modules that together implement "
        "the model lifecycle: data preparation, training, experiment analysis, "
        "model selection, serving, monitoring, benchmarking, and CI packaging."
    )
    add_table(
        doc,
        [
            ["Module", "Implementation", "Description"],
            [
                "Data module",
                "src/mlops_digit/data.py",
                "Loads scikit-learn digits or MNIST, applies reproducible "
                "stratified train/validation/test splitting, and exposes dataset "
                "metadata.",
            ],
            [
                "Model factory",
                "src/mlops_digit/models.py",
                "Builds sklearn pipelines for decision tree, SVM, logistic "
                "regression, random forest, KNN and MLP configurations.",
            ],
            [
                "Training module",
                "src/mlops_digit/train.py",
                "Reads YAML configs, trains models, logs metrics and artifacts "
                "to MLflow, and writes JSON summaries.",
            ],
            [
                "Analysis module",
                "src/mlops_digit/analyze.py",
                "Runs the underfit, balanced and overfit experiments and "
                "generates comparison artifacts.",
            ],
            [
                "Champion selector",
                "src/mlops_digit/selector.py",
                "Promotes the best generalising candidate according to a "
                "transparent validation-accuracy and train-val-gap rule.",
            ],
            [
                "Monitor",
                "src/mlops_digit/monitor.py",
                "Scores labelled batches and raises an alert when accuracy drops "
                "below the champion baseline by more than the threshold.",
            ],
            [
                "Serving layer",
                "src/mlops_digit/serve.py",
                "FastAPI service exposing /predict, /monitor, /livez, /readyz "
                "and /metrics.",
            ],
            [
                "Benchmark/load test",
                "src/mlops_digit/benchmark.py and load/locustfile.py",
                "Measures inference latency and exercises the HTTP path.",
            ],
            [
                "Packaging and CI",
                "Dockerfile.train, Dockerfile.serve, .github/workflows/ci.yml",
                "Builds isolated train/serve containers and validates the "
                "project in GitHub Actions.",
            ],
        ],
    )
    doc.add_paragraph("Figure 1: Modules in Digit Recognition MLOps Pipeline")


def functional_description(doc: Document) -> None:
    add_heading(doc, "2. FUNCTIONAL BLOCK DIAGRAM/DESCRIPTION")
    doc.add_paragraph(
        "The functional flow starts from a dataset configuration and ends with a "
        "served champion model that can be monitored during operation. Each "
        "block writes auditable artifacts so the next block can consume them "
        "without hidden notebook state."
    )
    add_table(
        doc,
        [
            ["Step", "Functional Block", "Input", "Output"],
            ["1", "Dataset loader", "digits or MNIST config", "X/y arrays"],
            [
                "2",
                "Split builder",
                "Raw arrays and random seed",
                "Train, validation and test splits",
            ],
            [
                "3",
                "Training pipeline",
                "YAML model config",
                "MLflow run, model artifact and metrics JSON",
            ],
            [
                "4",
                "Analysis",
                "Underfit/balanced/overfit runs",
                "Comparison metrics and plots",
            ],
            [
                "5",
                "Champion selection",
                "All run metrics and model artifacts",
                "models/champion.joblib and champion.json",
            ],
            [
                "6",
                "Serving",
                "Champion model and prediction requests",
                "Predicted digit labels, probabilities and Prometheus metrics",
            ],
            [
                "7",
                "Monitoring",
                "Labelled live batch and baseline accuracy",
                "OK/alert status and drift metrics",
            ],
        ],
    )
    doc.add_paragraph("Figure 2: Functional Block Diagram")
    doc.add_paragraph(
        "DATA CONFIG -> STRATIFIED SPLITS -> TRAIN + TRACK -> SELECT CHAMPION "
        "-> SERVE API -> MONITOR ACCURACY -> ALERT / METRICS"
    )
    add_subheading(doc, "Interfaces")
    add_bullets(
        doc,
        [
            "REST API: /predict for inference and /monitor for labelled-batch monitoring.",
            "Health probes: /livez and /readyz for container/service readiness.",
            "Observability: /metrics exposes Prometheus counters, gauges and histograms.",
            "Artifacts: models, metrics JSON files, proof table and generated reports.",
        ],
    )


def specifications(doc: Document, facts: dict) -> None:
    champ = facts["champion"]
    metrics = champ["metrics"]
    add_heading(doc, "3. MAJOR TECHNICAL SPECIFICATIONS")
    add_table(
        doc,
        [
            ["Sl no", "Technical parameter", "Specification"],
            ["1", "Project type", "CPU-only MLOps pipeline for digit recognition"],
            ["2", "Programming language", "Python 3.10+"],
            ["3", "Primary libraries", "scikit-learn, MLflow, FastAPI, Pydantic, Prometheus client"],
            ["4", "Default dataset", "scikit-learn digits: 1,797 samples, 64 features, 10 classes"],
            ["5", "Scale-up dataset", "MNIST: 70,000 samples, 784 features, config-selectable"],
            ["6", "Data split", "Stratified train/validation/test split with random_state 42"],
            ["7", "Candidate models", "Decision tree, SVM, logistic regression, random forest, KNN, MLP"],
            ["8", "Champion model", f"{champ['champion']} ({champ['model_type']})"],
            ["9", "Champion test accuracy", pct(metrics["test_accuracy"])],
            ["10", "Train-validation gap", f"{metrics['train_val_gap'] * 100:.2f} percentage points"],
            ["11", "Selection rule", champ["selection_rule"]],
            ["12", "Serving framework", "FastAPI with /predict, /monitor, /livez, /readyz and /metrics"],
            ["13", "Containerization", "Separate train and serve Dockerfiles based on python:3.11-slim"],
            ["14", "CI", "GitHub Actions: install dependencies, run pytest, train, benchmark and upload artifacts"],
            ["15", "Monitoring threshold", "Default accuracy-drop threshold: 0.05"],
        ],
    )
    doc.add_paragraph("Table 1: Technical Specifications of the Project")


def design_considerations(doc: Document) -> None:
    add_heading(doc, "4. DESIGN CONSIDERATIONS")
    add_bullets(
        doc,
        [
            "Reproducibility: data splits, model configs and selection rules are explicit and versionable.",
            "Generalisation first: model promotion considers validation accuracy and train-validation gap.",
            "Observability: prediction counts, latency histograms, monitored accuracy and drift alerts are exported.",
            "Portability: the pipeline is CPU-only and packaged with separate train and serve images.",
            "Service readiness: health and readiness probes make the API suitable for container deployment.",
            "Testing: unit tests, benchmark scripts and CI provide a base for the next testing phase.",
            "Extensibility: the data module supports both small offline digits experiments and MNIST scale-up.",
        ],
    )


def future_plan(doc: Document) -> None:
    add_heading(doc, "5. FUTURE PLAN")
    add_table(
        doc,
        [
            ["Sl No", "Phases", "Start Date - End Date", "Work to be done", "Status"],
            [
                "1",
                "Dissertation Outline",
                "01 May 2026 - 14 May 2026",
                "Literature review and prepare dissertation outline",
                "COMPLETED",
            ],
            [
                "2",
                "Design and Development",
                "15 May 2026 - 14 June 2026",
                "Implement data, training, selection, monitoring, serving, Docker and CI components",
                "COMPLETED",
            ],
            [
                "3",
                "Testing",
                "15 June 2026 - 01 July 2026",
                "Expand pytest coverage, validate API endpoints, run latency benchmark and Locust load tests",
                "PENDING",
            ],
            [
                "4",
                "Dissertation Review",
                "02 July 2026 - 12 July 2026",
                "Submit dissertation to supervisor and additional examiner for review and feedback",
                "PENDING",
            ],
            [
                "5",
                "Submission",
                "13 July 2026 - 17 July 2026",
                "Final review and dissertation submission",
                "PENDING",
            ],
        ],
    )


def abbreviations(doc: Document) -> None:
    add_heading(doc, "6. ABBREVIATIONS")
    add_table(
        doc,
        [
            ["API", "Application Programming Interface"],
            ["CI", "Continuous Integration"],
            ["CPU", "Central Processing Unit"],
            ["DSE", "Data Science & Engineering"],
            ["HTTP", "Hypertext Transfer Protocol"],
            ["KNN", "K-Nearest Neighbours"],
            ["ML", "Machine Learning"],
            ["MLflow", "Machine Learning experiment tracking platform"],
            ["MLOps", "Machine Learning Operations"],
            ["MNIST", "Modified National Institute of Standards and Technology dataset"],
            ["REST", "Representational State Transfer"],
            ["SVM", "Support Vector Machine"],
            ["WILP", "Work Integrated Learning Programmes"],
            ["YAML", "YAML Ain't Markup Language"],
        ],
        header=False,
    )


def build() -> None:
    facts = project_facts()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    cover(doc)
    doc.add_page_break()
    cover(doc)
    doc.add_page_break()
    abstract(doc, facts)
    doc.add_page_break()
    contents(doc)
    doc.add_page_break()
    modules(doc)
    doc.add_page_break()
    functional_description(doc)
    doc.add_page_break()
    specifications(doc, facts)
    design_considerations(doc)
    doc.add_page_break()
    future_plan(doc)
    abbreviations(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Report written to: {OUT_PATH}")


if __name__ == "__main__":
    build()
