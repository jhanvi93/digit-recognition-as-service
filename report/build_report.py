"""Generate the dissertation/project abstract .docx.

The layout faithfully follows the BITS WILP abstract template
(2024da04367-dissertation-abstract.docx) with the standard eight numbered
sections, a Plan of Work table, a Supervisor/Examiner particulars table, a
student-details table and a signature table.

Usage:
    python report/build_report.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = PROJECT_ROOT / "report" / "2024da04367-digit-recognition-abstract.docx"

TITLE = (
    "MLOps Pipeline for a Digit Recognition Service: Training, Experiment "
    "Tracking, Containerisation and Kubernetes Deployment"
)


# ----------------------------------------------------------------------------
# Low-level helpers
# ----------------------------------------------------------------------------
def center(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_centered(doc, label: str, value: str | None = None, bold_label: bool = True,
                 size: int | None = None) -> None:
    p = doc.add_paragraph()
    center(p)
    run = p.add_run(label)
    run.bold = bold_label
    if size:
        run.font.size = Pt(size)
    if value is not None:
        r2 = p.add_run(value)
        r2.bold = False
        if size:
            r2.font.size = Pt(size)


def add_heading(doc, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def add_bullets(doc, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it)


def set_table_borders(table) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl.tblPr.append(borders)


def add_table(doc, rows: list[list[str]], header: bool = True):
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            text = row[j] if j < len(row) else ""
            cell.text = text
            if header and i == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    return table


# ----------------------------------------------------------------------------
# Document
# ----------------------------------------------------------------------------
def build() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"

    # --- Title block ---------------------------------------------------------
    add_centered(doc, "Dissertation / Project / Project Work", size=18)
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    center(p)
    r = p.add_run("Dissertation / Project / Project Work Title:")
    r.bold = True
    p.add_run("\n\u201c" + TITLE + "\u201d")

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered(doc, "Course No.: ", "S2-25 DISSERTATION-NSP4")
    add_centered(doc, "Course Title: ", "Dissertation/Project (Merged Course-NSP4)")
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered(doc, "Dissertation / Project / Project Work Done by:")
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered(doc, "Student Name: ", "Jhanvi Titoria")
    add_centered(doc, "BITS ID: ", "2024da04367@wilp.bits-pilani.ac.in")
    add_centered(doc, "Degree Program: ", "M.Tech. Data Science & Engineering")
    add_centered(doc, "Research Area: ", "Data Science & Engineering applications")
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered(doc, "Dissertation / Project Work carried out at:")
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI")
    add_centered(doc, "VIDYA VIHAR, PILANI, RAJASTHAN - 333031.")
    add_centered(doc, "May-June 2026")
    doc.add_paragraph()

    # --- 1. Broad Area of Work ----------------------------------------------
    add_heading(doc, "1. Broad Area of Work")
    doc.add_paragraph(
        "Machine Learning Operations (MLOps) has emerged as a critical discipline "
        "at the intersection of software engineering, DevOps, and data science. "
        "The broad scope of this project covers the end-to-end lifecycle of a "
        "machine learning service for handwritten digit recognition, from model "
        "training and experiment tracking through containerization, orchestration, "
        "latency benchmarking, and continuous integration."
    )
    add_bullets(doc, [
        "Applied Machine Learning (Computer Vision): Ten-class handwritten digit "
        "classification on the scikit-learn 8x8 digits dataset (1,797 samples, "
        "64 features), running fully on CPU.",
        "Experiment Tracking & Model Governance: MLflow logging and comparison of "
        "three contrasted runs (underfit, balanced, overfit) with train/validation "
        "gap analysis.",
        "Model Packaging & Serving: joblib pipeline artifacts; FastAPI REST API "
        "with Pydantic request validation.",
        "Containerization: Separate multi-stage Docker images (Dockerfile.train, "
        "Dockerfile.serve).",
        "Cloud-Native Orchestration: Kubernetes-ready service with liveness and "
        "readiness probes for deployment and autoscaling.",
        "Observability & Reliability: /livez, /readyz, Prometheus /metrics; "
        "automated pytest suites.",
        "CI/CD & Performance Engineering: GitHub Actions; Locust load tests and an "
        "inference latency benchmark.",
    ])

    # --- 2. Background -------------------------------------------------------
    add_heading(doc, "2. Background")
    doc.add_paragraph(
        "Handwritten digit recognition is a canonical, well-understood "
        "classification problem, which makes it an ideal payload for "
        "demonstrating operational machine-learning infrastructure without the "
        "distraction of an exotic model. The engineering challenge addressed "
        "here is not the accuracy of a single classifier but the discipline of "
        "shipping, comparing, serving and monitoring models reliably."
    )
    doc.add_paragraph(
        "A large share of machine learning models never leaves the notebook. "
        "MLOps bridges the gap with reproducible training, comparable experiments, "
        "automated deployment, and observable services. This project deliberately "
        "trains three contrasted models on identical, stratified data splits to "
        "make the bias\u2013variance trade-off concrete: a high-bias underfit model, "
        "a well-regularised balanced model, and a high-variance overfit model. "
        "Each run is tracked in MLflow, packaged for serving, and characterised for "
        "inference latency."
    )

    # --- 3. Objectives -------------------------------------------------------
    add_heading(doc, "3. Objectives")
    add_numbered(doc, [
        "1. Understand the MLOps landscape and the engineering concerns that "
        "separate notebook models from production services.",
        "2. Implement reproducible, stratified train/validation/test splits of the "
        "digits dataset via mlops_digit.data.",
        "3. Implement config-driven training (mlops_digit.train) with YAML configs "
        "(underfit, balanced, overfit) and MLflow logging.",
        "4. Train and compare three contrasted runs; analyse the train vs "
        "validation accuracy gap to diagnose under- and overfitting.",
        "5. Build mlops_digit.serve (FastAPI) exposing /predict, /livez, /readyz "
        "and /metrics.",
        "6. Containerise training and serving as separate multi-stage Docker images.",
        "7. Prepare the service for Kubernetes deployment with health probes and "
        "resource limits.",
        "8. Add pytest suites, a Locust load test, an inference latency benchmark, "
        "and a GitHub Actions CI workflow.",
        "9. Validate end-to-end: local train \u2192 analyse \u2192 serve \u2192 benchmark, "
        "with health and metrics endpoints verified.",
    ])

    # --- 4. Scope of Work ----------------------------------------------------
    add_heading(doc, "4. Scope of Work")
    add_bullets(doc, [
        "Ten-class handwritten digit classification on the scikit-learn digits "
        "dataset with reproducible stratified splits.",
        "Three-model experimental arc (underfit / balanced / overfit) with MLflow "
        "tracking and automated under/overfitting diagnosis.",
        "FastAPI serving, Docker, Kubernetes readiness, automated tests, CI, and "
        "latency characterization.",
        "Dissertation chapters: problem, related work, methodology, experiments, "
        "results, conclusions.",
    ])

    # --- 5. Plan of Work -----------------------------------------------------
    add_heading(doc, "5. Plan of Work")
    add_table(doc, [
        ["Phases", "Start Date \u2013 End Date", "Work to be done"],
        ["Dissertation Outline", "01 May 2026 \u2013 14 May 2026",
         "Literature review on MLOps, model generalisation, containerization and "
         "cloud-native ML serving; prepare dissertation outline"],
        ["Design & Development", "15 May 2026 \u2013 14 June 2026",
         "Implement mlops_digit.data, mlops_digit.train, mlops_digit.serve, "
         "Dockerfiles and Kubernetes manifests"],
        ["Testing", "15 June 2026 \u2013 01 July 2026",
         "pytest suites, GitHub Actions CI; validate /livez, /readyz, /predict, "
         "/metrics; latency benchmark and Locust load test"],
        ["Dissertation Review", "02 July 2026 \u2013 12 July 2026",
         "Submit dissertation to Supervisor and Additional Examiner for review and "
         "feedback; incorporate revisions"],
        ["Submission", "13 July 2026 \u2013 17 July 2026",
         "Final review and submission of dissertation"],
    ])

    # --- 6. Literature References --------------------------------------------
    add_heading(doc, "6. Literature References")
    add_numbered(doc, [
        "1. D. Sculley et al. Hidden technical debt in machine learning systems. "
        "NeurIPS, 2015.",
        "2. D. Kreuzberger, N. K\u00fchl, S. Hirschl. Machine Learning Operations "
        "(MLOps): Overview, Definition, and Architecture. IEEE Access, 2023.",
        "3. M. Zaharia et al. Accelerating the Machine Learning Lifecycle with "
        "MLflow. IEEE Data Engineering Bulletin, 2018.",
        "4. B. Burns et al. Borg, Omega, and Kubernetes. Communications of the ACM, "
        "2016.",
        "5. S. Amershi et al. Software Engineering for Machine Learning: A Case "
        "Study. ICSE-SEIP, 2019.",
        "6. F. Pedregosa et al. Scikit-learn: Machine Learning in Python. JMLR, "
        "2011.",
        "7. C. M. Bishop. Pattern Recognition and Machine Learning. Springer, 2006.",
        "8. T. Hastie, R. Tibshirani, J. Friedman. The Elements of Statistical "
        "Learning. Springer, 2009.",
        "9. B. Brazil. Prometheus: Up & Running. O'Reilly, 2018.",
    ])

    # --- 7. Particulars of the Supervisor and Examiner -----------------------
    add_heading(doc, "7. Particulars of the Supervisor and Examiner")
    add_table(doc, [
        ["", "Supervisor", "Additional Examiner"],
        ["Name", "Abhishek Kumar", "Abhishek Kumar Jain"],
        ["Qualification",
         "BE (Electronics) from North Maharashtra University, Jalgaon; Executive "
         "Post Graduate Program in Business Management (Marketing), SIES College of "
         "Management Studies, Mumbai",
         "Executive Post Graduate Diploma in Business Management, Symbiosis "
         "International University, Pune; B.Tech (Electronics), Institute of "
         "Technology, Banaras Hindu University, Varanasi"],
        ["Designation", "R&D Specialist", "R&D Architect"],
        ["Employing Organization and Location",
         "Nokia Solutions and Networks, IN/Bangalore/Manyata E2",
         "Nokia Solutions and Networks, IN/Bangalore/Manyata E2"],
        ["Phone No. (with STD Code)", "+919620092234", "+919620048791"],
        ["Email Address", "abhishek.5.kumar@nokia.com", "abhishek.k.jain@nokia.com"],
    ])

    # --- 8. Remarks of the Supervisor ----------------------------------------
    add_heading(doc, "8. Remarks of the Supervisor")
    doc.add_paragraph(
        "The proposed study and project are well-aligned with current industrial "
        "practice in applied machine learning and DevOps. The student has scoped "
        "the work appropriately for a master's dissertation, using a well-understood "
        "digit-recognition task as the payload so that the emphasis falls on "
        "operational engineering. The project goes beyond model building to cover "
        "MLflow experiment tracking, separate training and serving containers, "
        "Kubernetes-ready health probes, Prometheus metrics, automated testing, "
        "latency benchmarking, and continuous integration. The three-run "
        "experimental design (underfit, balanced, overfit) provides a clear "
        "narrative for generalization and operational model selection. Deliverables "
        "are concrete and end-to-end verifiable. After discussion with the student, "
        "I am satisfied that the required engineering discipline and implementation "
        "skills are in place. I approve the project as supervisor."
    )

    # --- Footer block + tables ----------------------------------------------
    doc.add_paragraph()
    add_centered(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI")
    add_centered(doc, "WORK INTEGRATED LEARNING PROGRAMMES (WILP) DIVISION")
    add_centered(doc, "SECOND SEMESTER OF ACADEMIC YEAR 2025-2026")
    doc.add_paragraph()
    add_centered(doc, "S2-25 DISSERTATION-NSP4: Dissertation/Project (Merged Course-NSP4)")
    doc.add_paragraph()

    add_table(doc, [
        ["STUDENT ID No.", "2024DA04367"],
        ["NAME OF THE STUDENT", "Jhanvi Titoria"],
        ["STUDENT'S EMAIL ADDRESS", "2024da04367@wilp.bits-pilani.ac.in"],
        ["STUDENT'S EMPLOYING ORGANIZATION & LOCATION",
         "Nokia Solutions and Networks, IN/Bangalore/Manyata E2"],
        ["SUPERVISOR'S NAME", "Abhishek Kumar"],
        ["SUPERVISOR'S EMPLOYING ORGANIZATION & LOCATION",
         "Nokia Solutions and Networks, IN/Bangalore/Manyata E2"],
        ["SUPERVISOR'S EMAIL ADDRESS", "abhishek.5.kumar@nokia.com"],
        ["ADDITIONAL EXAMINER'S NAME", "Abhishek Kumar Jain"],
        ["ADDITIONAL EXAMINER'S EMPLOYING ORGANIZATION & LOCATION",
         "Nokia Solutions and Networks, IN/Bangalore/Manyata E2"],
        ["ADDITIONAL EXAMINER'S EMAIL ADDRESS", "abhishek.k.jain@nokia.com"],
        ["DISSERTATION / PROJECT / PROJECT WORK TITLE",
         "MLOps Pipeline for a digit recognition service: Training, "
         "Containerisation and Kubernetes Deployment"],
    ], header=False)
    doc.add_paragraph()

    add_table(doc, [
        ["Signature of Student", "Signature of Supervisor", "Signature of Additional Examiner"],
        ["Name: Jhanvi Titoria", "Name: Abhishek Kumar", "Name: Abhishek Kumar Jain"],
    ], header=False)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Report written to: {OUT_PATH}")


if __name__ == "__main__":
    build()
